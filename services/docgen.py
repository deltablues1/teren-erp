"""Generator građevinskog dnevnika i knjige po NN 60/2024.

Generira .docx programski (bez vanjskog template-a) pa po želji konvertira u .pdf."""
from __future__ import annotations

import logging
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from config import GENERATED_DIR
from services import claude_parser, repository as repo

log = logging.getLogger(__name__)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def _add_kv_table(doc: Document, pairs: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(pairs):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v or "—"
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(11)


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Calibri")


def _render_dnevnik_dan(
    doc: Document, projekt: dict[str, Any], projekt_key: str, datum: str,
    dnevnik_zapisi: list[dict[str, Any]], materijali_zapisi: list[dict[str, Any]],
    radnici_map: dict[str, dict[str, Any]],
) -> None:
    """Renderira sekcije jednog dana dnevnika (datum, vrijeme, radna snaga,
    opis radova, materijali po strujnom krugu, strojevi, napomene, potpisi)."""
    vrijeme = repo.get_weather_za_datum(projekt_key, datum)
    narativ = claude_parser.summarize_dnevnik(
        dnevnik_zapisi, materijali_zapisi, projekt.get("naziv", ""), datum,
    )

    _add_heading(doc, f"Datum: {datum}", level=1)

    _add_heading(doc, "Vremenske prilike", level=2)
    if vrijeme:
        _add_kv_table(doc, [
            ("Minimalna temperatura:", f"{vrijeme.get('Min_temp', '—')} °C"),
            ("Maksimalna temperatura:", f"{vrijeme.get('Max_temp', '—')} °C"),
            ("Oborine:", f"{vrijeme.get('Oborine_mm', 0)} mm"),
            ("Opis:", str(vrijeme.get("Vrijeme_opis", ""))),
        ])
    else:
        doc.add_paragraph("Podaci o vremenu nisu zabilježeni za ovaj dan.")

    _add_heading(doc, "Radna snaga", level=2)
    prisutni_ids = {str(r.get("Telegram_ID")) for r in dnevnik_zapisi}
    sati_po_radniku: dict[str, float] = defaultdict(float)
    vrijeme_po_radniku: dict[str, str] = {}
    spomenuti_imena: set[str] = set()
    for r in dnevnik_zapisi:
        tid = str(r.get("Telegram_ID"))
        try:
            sati_po_radniku[tid] += float(r.get("Sati") or 0)
        except (TypeError, ValueError):
            pass
        vr = str(r.get("Vrijeme_rada") or "")
        if vr and tid not in vrijeme_po_radniku:
            vrijeme_po_radniku[tid] = vr
        raw = str(r.get("Radnici_spomenuti") or "").strip()
        if raw:
            for ime in [x.strip() for x in raw.split(",") if x.strip()]:
                spomenuti_imena.add(ime)

    poznata_imena_pošiljatelja = {
        (radnici_map.get(tid, {}).get("Ime") or "").strip()
        for tid in prisutni_ids
    }
    spomenuti_dodatni = sorted(
        n for n in spomenuti_imena if n not in poznata_imena_pošiljatelja
    )

    if prisutni_ids or spomenuti_dodatni:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Ime"
        hdr[1].text = "Kvalifikacija"
        hdr[2].text = "Sati"
        hdr[3].text = "Vrijeme rada"
        for tid in sorted(prisutni_ids):
            r = radnici_map.get(tid)
            ime = (r.get("Ime") if r else "") or f"ID {tid}"
            kval = (r.get("Kvalifikacija") if r else "") or "Nespecificirano"
            sati = sati_po_radniku.get(tid, 0)
            row = table.add_row().cells
            row[0].text = ime
            row[1].text = kval
            row[2].text = f"{sati:g}" if sati else "—"
            row[3].text = vrijeme_po_radniku.get(tid, "—")
        for ime in spomenuti_dodatni:
            row = table.add_row().cells
            row[0].text = ime
            row[1].text = "(spomenut, ne-registriran)"
            row[2].text = "—"
            row[3].text = "—"

        ukupno = len(prisutni_ids) + len(spomenuti_dodatni)
        p_total = doc.add_paragraph()
        run = p_total.add_run(f"Ukupno ljudi na gradilištu: {ukupno}")
        run.italic = True
    else:
        doc.add_paragraph("Tog dana nije zabilježena prisutnost radnika.")

    _add_heading(doc, "Opis izvedenih radova", level=2)
    if narativ:
        for para in narativ.split("\n\n"):
            doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph("Nema zabilježenih radova.")

    # Napomena: ugrađeni materijali (količine po strujnim krugovima) NE idu u
    # dnevnik — to je obračunski podatak za građevinsku knjigu.

    _add_heading(doc, "Strojevi i oprema na gradilištu", level=2)
    doc.add_paragraph("(nadopuniti ručno)")

    _add_heading(doc, "Posebne napomene / incidenti", level=2)
    problemi_zbirno: list[str] = []
    for r in dnevnik_zapisi:
        raw = str(r.get("Problemi") or "").strip()
        if raw:
            for chunk in raw.split("|"):
                chunk = chunk.strip()
                if chunk:
                    problemi_zbirno.append(f"{r.get('Radnik', '?')}: {chunk}")
    if problemi_zbirno:
        for p in problemi_zbirno:
            doc.add_paragraph(p, style="List Bullet")
    else:
        doc.add_paragraph("Nema posebnih napomena.")

    _add_heading(doc, "Potpisi", level=2)
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.rows[0].cells[0].text = "Odgovorna osoba izvođača:"
    sig_table.rows[0].cells[1].text = "Nadzorni inženjer:"
    sig_table.rows[1].cells[0].text = "\n\n______________________\n"
    sig_table.rows[1].cells[1].text = "\n\n______________________\n"


def generate_dnevnik(projekt_key: str, od: str, do: str | None = None) -> Path:
    """Generira .docx građevinski dnevnik za projekt u rasponu [od, do].

    do=None → samo dan `od`. Renderira jednu stranicu po danu koji ima unose
    (dnevnik ili materijal); materijali su grupirani po strujnom krugu.
    """
    projekt = repo.get_projekt(projekt_key)
    if not projekt:
        raise ValueError(f"Projekt '{projekt_key}' ne postoji.")
    do = do or od

    dnevnik_all = repo.get_dnevnik_period(projekt_key, od=od, do=do)
    materijali_all = repo.get_materijali_period(projekt_key, od=od, do=do)
    radnici_lista = repo.list_radnici(projekt_key)
    radnici_map = {str(r.get("Telegram_ID")): r for r in radnici_lista}

    # dani s podacima u rasponu (po datumu)
    dani = sorted(
        {str(r.get("Datum")) for r in dnevnik_all}
        | {str(m.get("Datum")) for m in materijali_all}
    )

    doc = Document()
    _set_default_font(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GRAĐEVINSKI DNEVNIK")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(
        "sukladno Pravilniku o načinu provedbe stručnog nadzora građenja, "
        "vođenju građevinskog dnevnika (NN 60/2024)"
    )
    sr.italic = True
    sr.font.size = Pt(9)

    if not od and not do:
        razdoblje, raspon_dio = "cijeli projekt", "cijeli"
    elif od == do:
        razdoblje, raspon_dio = od, od
    else:
        razdoblje, raspon_dio = f"{od} – {do}", f"{od}_{do}"
    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.add_run(f"Razdoblje: {razdoblje}").italic = True

    doc.add_paragraph()

    _add_heading(doc, "Opći podaci o gradilištu", level=1)
    _add_kv_table(doc, [
        ("Naziv objekta:", projekt.get("naziv", "")),
        ("Adresa gradilišta:", projekt.get("adresa", "")),
        ("Investitor:", projekt.get("investitor", "")),
        ("Izvođač:", projekt.get("izvodac", "")),
        ("Nadzorni inženjer:", projekt.get("nadzorni", "")),
        ("Broj građevinske dozvole:", projekt.get("broj_dozvole", "")),
    ])

    if not dani:
        doc.add_paragraph()
        doc.add_paragraph("U odabranom razdoblju nema zabilježenih unosa.")
    for i, datum in enumerate(dani):
        doc.add_page_break()
        dnevnik_dan = [r for r in dnevnik_all if str(r.get("Datum")) == datum]
        materijali_dan = [m for m in materijali_all if str(m.get("Datum")) == datum]
        _render_dnevnik_dan(
            doc, projekt, projekt_key, datum, dnevnik_dan, materijali_dan, radnici_map,
        )

    safe_naziv = projekt.get("naziv", projekt_key).replace(" ", "_")
    out_path = GENERATED_DIR / f"dnevnik_{safe_naziv}_{raspon_dio}.docx"
    doc.save(out_path)
    log.info("Generiran dnevnik: %s (%d dana)", out_path, len(dani))
    return out_path


def _num(v: Any) -> float:
    try:
        return float(v or 0)
    except (TypeError, ValueError):
        return 0.0


def _last_day_of_month(year: int, month: int) -> int:
    from calendar import monthrange
    return monthrange(year, month)[1]


def _resolve_period(mjesec: str | None) -> tuple[str | None, str]:
    """mjesec='YYYY-MM' → (YYYY-MM-01, YYYY-MM-LAST). None → (None, danas)."""
    if not mjesec:
        return None, datetime.now().strftime("%Y-%m-%d")
    try:
        year, month = (int(x) for x in mjesec.split("-")[:2])
    except (ValueError, IndexError):
        raise ValueError(f"Neispravan mjesec '{mjesec}', očekujem YYYY-MM")
    od = f"{year:04d}-{month:02d}-01"
    do = f"{year:04d}-{month:02d}-{_last_day_of_month(year, month):02d}"
    return od, do


def _add_knjiga_page(
    doc: Document, projekt: dict[str, Any], stavka: dict[str, Any],
    izvr_mjesecno: float, izvr_kumulativ: float,
    stranica: int, situacija_broj: int | None,
    krug_razrada: dict[str, float] | None = None,
) -> None:
    """Doda jednu stranicu knjige za jednu stavku troškovnika (obračunski format)."""
    sifra = str(stavka.get("Šifra", "")).strip()
    opis_stavke = str(stavka.get("Opis stavke", "")).strip()
    pozicija = str(stavka.get("Pozicija", "")).strip() or opis_stavke
    jm = str(stavka.get("JM", "")).strip()
    planirano = _num(stavka.get("Ugovorena količina"))
    cijena = _num(stavka.get("Jedinična cijena"))

    iznos_mjesecno = izvr_mjesecno * cijena
    iznos_kumulativ = izvr_kumulativ * cijena

    # 1) Zaglavlje stranice: GRAĐEVINA + stranica
    head = doc.add_table(rows=2, cols=4)
    head.style = "Light Grid Accent 1"
    head.rows[0].cells[0].text = "GRAĐEVINA:"
    head.rows[0].cells[1].text = projekt.get("naziv", "")
    head.rows[0].cells[2].text = "Stranica:"
    head.rows[0].cells[3].text = str(stranica)
    head.rows[1].cells[0].text = "OPIS RADOVA:"
    merged = head.rows[1].cells[1].merge(head.rows[1].cells[3])
    merged.text = pozicija

    doc.add_paragraph()  # spacer

    # 2) Tablica obračuna - jedna stavka
    tbl = doc.add_table(rows=2, cols=8)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Redni broj"
    hdr[1].text = "JM"
    hdr[2].text = "Ukupna kol. po troškovniku"
    hdr[3].text = "Jed. cijena"
    hdr[4].text = "Izvršeno mjesečno"
    hdr[5].text = "Izvršeno ukupno"
    hdr[6].text = "Iznos mjesečno"
    hdr[7].text = "Iznos ukupno"

    row = tbl.rows[1].cells
    row[0].text = sifra
    row[1].text = jm
    row[2].text = f"{planirano:g}" if planirano else "—"
    row[3].text = f"{cijena:g}" if cijena else "—"
    row[4].text = f"{izvr_mjesecno:g}"
    row[5].text = f"{izvr_kumulativ:g}"
    row[6].text = f"{iznos_mjesecno:.2f}" if cijena else "—"
    row[7].text = f"{iznos_kumulativ:.2f}" if cijena else "—"

    doc.add_paragraph()

    # 3) Opis stavke (kratki) ako se razlikuje od pozicije
    if opis_stavke and opis_stavke != pozicija:
        p = doc.add_paragraph()
        p.add_run("Stavka: ").bold = True
        p.add_run(opis_stavke)

    # 4) UKUPNO za stavku + oznaka situacije
    if situacija_broj:
        p_sit = doc.add_paragraph()
        p_sit.add_run(f"{situacija_broj}. SITUACIJA: ").bold = True
        p_sit.add_run(f"{izvr_mjesecno:g} {jm}")

    p_uk = doc.add_paragraph()
    p_uk.add_run("UKUPNO IZVEDENO: ").bold = True
    p_uk.add_run(f"{izvr_kumulativ:g} {jm}  od planiranih  {planirano:g} {jm}")
    razlika = planirano - izvr_kumulativ
    if planirano and abs(razlika) > 1e-9:
        p_uk.add_run(f"   (preostalo: {razlika:g} {jm})")

    # 4b) Razrada izvedenog po STRUJNIM KRUGOVIMA (dokaznica mjera) — kad postoji
    if krug_razrada:
        doc.add_paragraph()
        p_r = doc.add_paragraph()
        p_r.add_run("Razrada po strujnim krugovima:").bold = True
        rt = doc.add_table(rows=1, cols=2)
        rt.style = "Light Grid Accent 1"
        rt.rows[0].cells[0].text = "Strujni krug"
        rt.rows[0].cells[1].text = f"Izvršeno ({jm})" if jm else "Izvršeno"
        # krugovi sortirano, '(bez kruga)' zadnji
        for krug in sorted(krug_razrada, key=lambda k: (k == "(bez kruga)", k)):
            row = rt.add_row().cells
            row[0].text = krug
            row[1].text = f"{krug_razrada[krug]:g}"

    doc.add_paragraph()

    # 5) Potpisi (po stranici knjige)
    sig = doc.add_table(rows=2, cols=2)
    sig.rows[0].cells[0].text = "IZVOĐAČ:"
    sig.rows[0].cells[1].text = "NADZORNI INŽENJER:"
    sig.rows[1].cells[0].text = "\n______________________\n"
    sig.rows[1].cells[1].text = "\n______________________\n"


def _add_rekapitulacija(
    doc: Document, projekt: dict[str, Any], troskovnik: list[dict[str, Any]],
    izvr_kumulativ: dict[str, float], izvr_mjesecno: dict[str, float],
    situacija_broj: int | None,
) -> None:
    """Sumarna stranica po sekcijama."""
    doc.add_page_break()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("REKAPITULACIJA")
    run.bold = True
    run.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Projekt: {projekt.get('naziv', '')}")
    if situacija_broj:
        sub.add_run(f"   |   {situacija_broj}. SITUACIJA")

    doc.add_paragraph()

    # zbroj po sekcijama
    iznos_mj_po_sekc: dict[str, float] = defaultdict(float)
    iznos_uk_po_sekc: dict[str, float] = defaultdict(float)
    plan_po_sekc: dict[str, float] = defaultdict(float)
    for s in troskovnik:
        sekc = str(s.get("Sekcija", "")).strip() or "OSTALO"
        sifra = str(s.get("Šifra", "")).strip()
        cijena = _num(s.get("Jedinična cijena"))
        plan = _num(s.get("Ugovorena količina"))
        iznos_mj_po_sekc[sekc] += izvr_mjesecno.get(sifra, 0) * cijena
        iznos_uk_po_sekc[sekc] += izvr_kumulativ.get(sifra, 0) * cijena
        plan_po_sekc[sekc] += plan * cijena

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Sekcija"
    hdr[1].text = "Ugovoreno (€)"
    hdr[2].text = "Mjesečno (€)"
    hdr[3].text = "Kumulativno (€)"

    total_plan = total_mj = total_uk = 0.0
    for sekc in plan_po_sekc:
        row = tbl.add_row().cells
        row[0].text = sekc
        row[1].text = f"{plan_po_sekc[sekc]:.2f}"
        row[2].text = f"{iznos_mj_po_sekc[sekc]:.2f}"
        row[3].text = f"{iznos_uk_po_sekc[sekc]:.2f}"
        total_plan += plan_po_sekc[sekc]
        total_mj += iznos_mj_po_sekc[sekc]
        total_uk += iznos_uk_po_sekc[sekc]

    row = tbl.add_row().cells
    row[0].text = "SVEUKUPNO"
    row[1].text = f"{total_plan:.2f}"
    row[2].text = f"{total_mj:.2f}"
    row[3].text = f"{total_uk:.2f}"
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def generate_knjiga(
    projekt_key: str,
    situacija_broj: int | None = None,
    mjesec: str | None = None,
) -> Path:
    """Generira građevinsku knjigu / obračunsku situaciju u standardnom obračunskom formatu.

    - situacija_broj: redni broj situacije (1, 2, ...) ili None za kumulativ
    - mjesec: 'YYYY-MM' za obračun konkretnog mjeseca, None = stanje na danas
    """
    projekt = repo.get_projekt(projekt_key)
    if not projekt:
        raise ValueError(f"Projekt '{projekt_key}' ne postoji.")

    troskovnik = repo.get_troskovnik(projekt_key)
    period_od, period_do = _resolve_period(mjesec)

    # svi materijali do kraja perioda → kumulativ
    svi_do = repo.get_materijali_period(projekt_key, do=period_do)
    izvr_kumulativ: dict[str, float] = defaultdict(float)
    for m in svi_do:
        sifra = str(m.get("Šifra_stavke", "")).strip()
        if not sifra:
            continue
        izvr_kumulativ[sifra] += _num(m.get("Količina"))

    # materijali samo u zadanom mjesecu → mjesečno
    izvr_mjesecno: dict[str, float] = defaultdict(float)
    if period_od:
        u_mjesecu = repo.get_materijali_period(
            projekt_key, od=period_od, do=period_do,
        )
        for m in u_mjesecu:
            sifra = str(m.get("Šifra_stavke", "")).strip()
            if not sifra:
                continue
            izvr_mjesecno[sifra] += _num(m.get("Količina"))
    else:
        izvr_mjesecno = izvr_kumulativ  # bez perioda - mjesečno = kumulativ

    # Eksplicitne veze materijal→troškovnička stavka (panel) imaju prednost pred
    # podudaranjem po šifri artikla — keširano po ŠIFRI troškovnika. Primijeni
    # samo ako postoji barem jedna veza (inače ostaje gornja logika).
    from services import db_backend
    veze_kum, veze_mj = db_backend.izvedeno_po_sifri(projekt_key, period_od, period_do)
    if veze_kum:
        izvr_kumulativ = veze_kum
        izvr_mjesecno = veze_mj if period_od else dict(veze_kum)

    # Ako postoje uvezene situacije (iz Excela), koristi snapshot kao izvor
    # izvedenih količina umjesto materijala radnika. Konkretna situacija →
    # kumulativ te situacije, "mjesečno" = razlika prema prethodnoj. Bez zadanog
    # broja (opća knjiga) → zadnja situacija, "mjesečno" = kumulativ.
    broj = situacija_broj or db_backend.zadnja_situacija_broj(projekt_key)
    if broj:
        snap, prev = db_backend.situacija_kumulativ(projekt_key, broj)
        if snap:
            izvr_kumulativ = snap
            if situacija_broj:
                izvr_mjesecno = {
                    sifra: kol - prev.get(sifra, 0.0) for sifra, kol in snap.items()
                }
            else:
                izvr_mjesecno = dict(snap)

    # generiraj dokument
    doc = Document()
    _set_default_font(doc)

    # naslovna
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GRAĐEVINSKA KNJIGA")
    run.bold = True
    run.font.size = Pt(20)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        f"{situacija_broj}. SITUACIJA" if situacija_broj else "Kumulativno stanje"
    )
    sr.font.size = Pt(14)

    _add_kv_table(doc, [
        ("Investitor:", projekt.get("investitor", "")),
        ("Izvođač:", projekt.get("izvodac", "")),
        ("Nadzorni inženjer:", projekt.get("nadzorni", "")),
        ("Adresa gradilišta:", projekt.get("adresa", "")),
        ("Razdoblje od:", period_od or "od početka"),
        ("Razdoblje do:", period_do),
    ])

    # razrada izvedenog po strujnim krugovima (iz povezanih materijala s terena)
    krug_po_sifri = db_backend.izvedeno_krug_po_sifri(projekt_key)

    # stranica po stavci
    for i, stavka in enumerate(troskovnik, start=1):
        doc.add_page_break()
        sifra = str(stavka.get("Šifra", "")).strip()
        _add_knjiga_page(
            doc, projekt, stavka,
            izvr_mjesecno.get(sifra, 0.0),
            izvr_kumulativ.get(sifra, 0.0),
            stranica=i,
            situacija_broj=situacija_broj,
            krug_razrada=krug_po_sifri.get(sifra),
        )

    _add_rekapitulacija(
        doc, projekt, troskovnik, izvr_kumulativ, izvr_mjesecno, situacija_broj,
    )

    safe_naziv = projekt.get("naziv", projekt_key).replace(" ", "_")
    sit_dio = f"_SIT{situacija_broj}" if situacija_broj else ""
    period_dio = f"_{mjesec}" if mjesec else f"_{datetime.now():%Y-%m-%d}"
    out_path = GENERATED_DIR / f"knjiga_{safe_naziv}{sit_dio}{period_dio}.docx"
    doc.save(out_path)
    log.info("Generirana knjiga: %s", out_path)
    return out_path


def generate_ponuda(ponuda: dict[str, Any]) -> Path:
    """Generira .docx ponudu iz dicta (services/ponude.get). Vrati putanju."""
    import config as cfg

    doc = Document()
    _set_default_font(doc)

    # zaglavlje firme
    firma_linije = [x for x in (
        cfg.FIRMA_NAZIV,
        cfg.FIRMA_ADRESA,
        f"OIB: {cfg.FIRMA_OIB}" if cfg.FIRMA_OIB else "",
        f"IBAN: {cfg.FIRMA_IBAN}" if cfg.FIRMA_IBAN else "",
        " · ".join(x for x in (cfg.FIRMA_TELEFON, cfg.FIRMA_EMAIL) if x),
    ) if x]
    for i, linija in enumerate(firma_linije):
        p = doc.add_paragraph()
        run = p.add_run(linija)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"PONUDA br. {ponuda['broj']}")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()
    _add_kv_table(doc, [
        ("Kupac:", ponuda.get("kupac_naziv", "")),
        ("Adresa:", ponuda.get("kupac_adresa", "")),
        ("OIB:", ponuda.get("kupac_oib", "")),
        ("Datum:", ponuda.get("datum", "")),
        ("Valjanost ponude:", f"{ponuda.get('valjanost_dana', 30)} dana"),
        ("Predmet:", ponuda.get("predmet", "")),
    ])

    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, naslov in enumerate(("Rb.", "Opis", "JM", "Količina", "Cijena (€)", "Iznos (€)")):
        hdr[i].text = naslov
    for i, st in enumerate(ponuda.get("stavke", []), start=1):
        row = tbl.add_row().cells
        row[0].text = str(i)
        row[1].text = str(st.get("opis", ""))
        row[2].text = str(st.get("jm", ""))
        row[3].text = f"{st.get('kolicina', 0):g}"
        row[4].text = f"{st['cijena']:.2f}" if st.get("cijena") is not None else "—"
        row[5].text = f"{st['iznos']:.2f}" if st.get("iznos") is not None else "—"

    doc.add_paragraph()
    from services.ponude import PDV_STOPA
    sume = doc.add_table(rows=3, cols=2)
    sume.style = "Light Grid Accent 1"
    sume.rows[0].cells[0].text = "Osnovica (€):"
    sume.rows[0].cells[1].text = f"{ponuda.get('osnovica', 0):.2f}"
    sume.rows[1].cells[0].text = f"PDV {PDV_STOPA:g}% (€):"
    sume.rows[1].cells[1].text = f"{ponuda.get('pdv', 0):.2f}"
    sume.rows[2].cells[0].text = "UKUPNO (€):"
    sume.rows[2].cells[1].text = f"{ponuda.get('ukupno', 0):.2f}"
    for p_ in sume.rows[2].cells[0].paragraphs + sume.rows[2].cells[1].paragraphs:
        for r in p_.runs:
            r.bold = True

    if ponuda.get("napomena"):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Napomena: ").bold = True
        p.add_run(ponuda["napomena"])

    doc.add_paragraph()
    doc.add_paragraph()
    potpis = doc.add_paragraph()
    potpis.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    potpis.add_run(f"{cfg.FIRMA_NAZIV or 'Ponuditelj'}\n\n______________________")

    out_path = GENERATED_DIR / f"ponuda_{ponuda['broj'].replace('/', '-')}.docx"
    doc.save(out_path)
    log.info("Generirana ponuda: %s", out_path)
    return out_path


def to_pdf(docx_path: Path) -> Path | None:
    """Konvertiraj .docx u .pdf preko docx2pdf (treba MS Word ili LibreOffice).
    Vrati None ako konverzija ne uspije - .docx je i dalje dostupan."""
    try:
        from docx2pdf import convert
        pdf_path = docx_path.with_suffix(".pdf")
        convert(str(docx_path), str(pdf_path))
        return pdf_path
    except Exception as e:
        log.warning("PDF konverzija nije uspjela (%s) - vraćam samo .docx", e)
        return None
